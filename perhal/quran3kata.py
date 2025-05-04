import json
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def format_paragraph(paragraph):
    run = paragraph.runs[0]
    run.font.name = 'Traditional Arabic'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Traditional Arabic')
    run.font.size = Pt(12)
    paragraph.paragraph_format.line_spacing = Pt(10)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

def get_three_words(word_list):
    return ' '.join([w['ar'] for w in word_list[:3]])

# Load JSON
with open('quran_data.json', encoding='utf-8') as f:
    quran_data = json.load(f)

# Ambil info halaman -> list ayat (3 kata) + juz & surat pertama
halaman_quran = {}

for halaman_str in quran_data:
    halaman = int(halaman_str)
    ayats = quran_data[halaman_str]

    three_words_lines = []
    juz_info = None
    surat_info = None

    for key in sorted(ayats.keys(), key=lambda x: (int(x.split('_')[0]), int(x.split('_')[1]))):
        ayat = ayats[key]
        words = ayat['words']
        if not juz_info:
            juz_info = ayat['juz']
        if not surat_info:
            surat_info = ayat['nama_surat']
        three_words_lines.append(get_three_words(words))
    
    halaman_quran[halaman] = {
        'juz': juz_info,
        'surat': surat_info,
        'lines': three_words_lines
    }

# Mulai dokumen
doc = Document()

i = 1
while i <= 604:
    # Ambil data halaman ganjil
    upper = halaman_quran.get(i)
    lower = halaman_quran.get(i + 1) if i + 1 <= 604 else None

    # Info header di atas tabel
    if upper:
        header_text = f"Juz {upper['juz']} | Hal {i} | {upper['surat']}"
    elif lower:
        header_text = f"Juz {lower['juz']} | Hal {i+1} | {lower['surat']}"
    else:
        header_text = f"Hal {i}–{i+1}"

    header_para = doc.add_paragraph(header_text)
    format_paragraph(header_para)

    # Buat tabel 2 baris
    table = doc.add_table(rows=2, cols=1)
    table.allow_autofit = False

    for row_offset, data in enumerate([upper, lower]):
        if not data:
            continue
        cell = table.rows[row_offset].cells[0]
        para = cell.add_paragraph(f"Juz {str(data['juz']).zfill(2)} - {i + row_offset}")
        format_paragraph(para)
        for line in data['lines']:
            para = cell.add_paragraph(line)
            format_paragraph(para)

    doc.add_page_break()
    i += 2

# Simpan dokumen
output_filename = "al_quran_3_kata_per_ayat_604_halaman.docx"
doc.save(output_filename)
print(f"✅ Berhasil disimpan sebagai: {output_filename}")

