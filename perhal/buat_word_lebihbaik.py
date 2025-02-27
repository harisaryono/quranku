import json
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
from docx.oxml import OxmlElement, ns

def set_paragraph_background(paragraph, color):
    """Mengatur warna latar belakang paragraf dalam Word menggunakan XML"""
    p = paragraph._element
    pPr = p.find(ns.qn('w:pPr'))
    if pPr is None:
        pPr = OxmlElement('w:pPr')
        p.append(pPr)
    
    shd = OxmlElement('w:shd')
    shd.set(ns.qn('w:val'), "clear")
    shd.set(ns.qn('w:color'), "auto")
    shd.set(ns.qn('w:fill'), color)
    pPr.append(shd)

c_par=['EAEAEA', 'F8F8F8']  # Warna paragraf abu-abu terang
c_word=['FFFFCC', 'FFFFFF']  # Kata: Kuning pastel & putih

def alternate_paragraph_color(index):
    """Mengembalikan warna latar belakang untuk paragraf bergantian"""
    return c_par[index % len(c_par)]

def alternate_word_color(index):
    """Mengembalikan warna latar belakang untuk kata bergantian"""
    return c_word[index % len(c_word)]

def get_text_color(bg_color):
    """Mengatur warna teks agar kontras dengan background"""
    if bg_color == 'FFFFCC':  # Jika latar belakang kuning pastel, teks merah
        return RGBColor(255, 0, 0)
    return RGBColor(0, 0, 0)  # Default teks hitam

def write_quran_to_word(data_file, output_file):
    with open(data_file, 'r', encoding='utf-8') as f:
        quran_data = json.load(f)

    document = Document()
    page_count = 0  

    for page, page_data in quran_data.items():
        #if page_count >= 10:
        #    break

        ayat_keys = list(page_data.keys())  
        nama_surat = page_data[ayat_keys[0]]['nama_surat']
        juz = page_data[ayat_keys[0]]['juz']

        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_title = paragraph.add_run(f"juz {juz}.: سورة {nama_surat}  :. hal {page}")
        run_title.font.name = 'Arial'
        run_title.font.size = Pt(10)
        run_title.bold = True

        for index, (ayat, ayat_data) in enumerate(page_data.items()):
            nomor_surat, nomor_ayat = map(int, ayat.split('_'))

            if nomor_ayat == 1:
                paragraph = document.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run_ayat = paragraph.add_run(f"{nomor_surat}. سورة {ayat_data['nama_surat']}")
                run_ayat.font.name = 'Arial'
                run_ayat.font.size = Pt(14)
                run_ayat.bold = True

            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            set_paragraph_background(paragraph, alternate_paragraph_color(index))

            run_nomor_surat = paragraph.add_run(f"{nomor_ayat}. ")
            run_nomor_surat.font.name = 'Arial'
            run_nomor_surat.font.size = Pt(12)
            run_nomor_surat.bold = True

            kata_index = 0
            for word in ayat_data['words']:
                bg_color = alternate_word_color(kata_index)
                text_color = get_text_color(bg_color)
                
                run_word = paragraph.add_run(f" {word['id']} ")  # Tambah spasi untuk padding
                run_word.font.name = 'Arial'
                run_word.font.size = Pt(12)
                run_word.font.color.rgb = text_color
                run_word.bold = (bg_color == 'FFFFCC')  # Jika kuning pastel, bold
                
                rPr = run_word._element.find(ns.qn('w:rPr'))
                if rPr is None:
                    rPr = OxmlElement('w:rPr')
                    run_word._element.append(rPr)

                shd = OxmlElement('w:shd')
                shd.set(ns.qn('w:val'), "clear")
                shd.set(ns.qn('w:fill'), bg_color)
                rPr.append(shd)
                
                kata_index += 1

        page_count += 1  
        document.add_page_break()

    document.save(output_file)

# Contoh penggunaan
write_quran_to_word('quran_data.json', 'quran_word.docx')
