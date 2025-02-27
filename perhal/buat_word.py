import json
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_COLOR
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml import ns  # Impor modul ns secara keseluruhan
from docx.oxml import OxmlElement


def alternate_color(index):
    colors = [WD_COLOR_INDEX.YELLOW, WD_COLOR_INDEX.BRIGHT_GREEN]  # Biru dan putih
    return colors[index % len(colors)]


def alternate_paragraph_color(index):
    colors = ['FAFAFA', 'EAF6FF']  # Putih dan abu-abu terang (Hex color codes)
    return colors[index % len(colors)]

def set_paragraph_background(paragraph, color):
    """Sets the background color of a paragraph."""
    shd = OxmlElement('w:shd')
    shd.set(ns.qn('w:val'), "clear")  # "clear", "solid", "nil"
    shd.set(ns.qn('w:color'), "auto")
    shd.set(ns.qn('w:fill'), color)  # Hex color code
    paragraph._p.append(shd)

def write_quran_to_word(data_file, output_file):
    with open(data_file, 'r', encoding='utf-8') as f:
        quran_data = json.load(f)

    document = Document()
    page_count = 0  # Hitung jumlah halaman yang diproses

    for page, page_data in quran_data.items():

        if page_count >= 10:  # Batasi hanya 10 halaman
           break
        
        document.add_section(WD_SECTION.NEW_PAGE)

        # Isi Halaman (Semua informasi sekarang di sini)
        ayat_keys = list(page_data.keys())  # Mengambil semua key ayat yang ada di halaman
        nama_surat = page_data[ayat_keys[0]]['nama_surat']
        juz = page_data[ayat_keys[0]]['juz']

        # Judul halaman (nama surat dan juz)
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_title = paragraph.add_run(f"{nama_surat} (Juz {juz}) - hal {page}")
        run_title.font.name = 'Arial'
        run_title.font.size = Pt(14)
        run_title.bold = True

        for ayat, ayat_data in page_data.items():
            nomor_surat, nomor_ayat = map(int, ayat.split('_'))  # Split kunci ayat

            if nomor_ayat == 1:  # Tulis heading hanya jika nomor ayat 1
                paragraph = document.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run_ayat = paragraph.add_run(f"{nomor_surat}. {ayat_data['nama_surat']}")
                run_ayat.font.name = 'Arial'
                run_ayat.font.size = Pt(14)
                run_ayat.bold = True
                print(f"{nomor_surat}. {ayat_data['nama_surat']}")

            paragraph = document.add_paragraph()
            paragraph_format = paragraph.paragraph_format
            paragraph_format.space_before = Pt(6) # mengatur jarak antar paragraph
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT #rata kiri

            set_paragraph_background(paragraph, alternate_paragraph_color(nomor_ayat-1)) #warna latar per paragraph

            run_nomor_surat = paragraph.add_run(f"{nomor_ayat}. ") #menambahkan nomor ayat di awal ayat
            run_nomor_surat.font.name = 'Arial'
            run_nomor_surat.font.size = Pt(12)
            run_nomor_surat.bold = True

            kata_index = 0
            for word in ayat_data['words']:
                run_word = paragraph.add_run(f"{word['id']} ") #spasi setelah kata
                run_word.font.name = 'Arial'
                run_word.font.size = Pt(12)
                run_word.font.highlight_color = alternate_color(kata_index)  # Warna alternatif
                kata_index += 1

        page_count += 1  # Tambahkan hitungan halaman setelah diproses
    document.save(output_file)

# Contoh penggunaan
write_quran_to_word('quran_data.json', 'quran_word.docx')